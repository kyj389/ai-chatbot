from docx import Document
import random
import os
import re

# 📂 문서 저장 경로
os.makedirs("./docs", exist_ok=True)

# 샘플 값 목록
companies = ["신한은행", "신한카드", "신한라이프", "신한투자증권"]
keywords = ["AI", "플랫폼", "인프라", "UX", "오픈금융", "키오스크", "UI", "서버", "통합", "스마트"]
locations = ["서울"]
durations = [f"{i}개월" for i in range(6, 25)]  # 6개월부터 24개월까지 1씩 증가
prices = [10_000_000,
          20_000_000,
          30_000_000,
          40_000_000, 
          50_000_000, 
          100_000_000, 
          200_000_000, 
          300_000_000, 
          400_000_000, 
          500_000_000, 
          1_000_000_000, 
          2_000_000_000, 
          3_000_000_000, 
          4_000_000_000, 
          5_000_000_000, 
          6_000_000_000, 
          7_000_000_000, 
          8_000_000_000, 
          9_000_000_000, 
          10_000_000_000]
descriptions = [
    "신규 플랫폼 구축 및 AI 연계 서비스 도입",
    "기존 시스템 고도화 및 UI/UX 개선",
    "통합 인증 시스템 도입",
    "모바일 앱 재구축 및 데이터 연계 개발",
    "스마트 키오스크 신규 배치",
    "디지털 뱅킹 서비스 개선 및 고객 경험 향상",
    "AI 기반 리스크 관리 시스템 도입",
    "블록체인 기술을 활용한 거래 시스템 구축",
    "모바일 뱅킹 앱 기능 추가 및 보안 강화",
    "고객 데이터 분석을 통한 맞춤형 서비스 제공"
]

# 🔄 랜덤 mock 문서 100개 생성
for i in range(1, 101):
    doc = Document()
    company = random.choice(companies)
    keyword = random.choice(keywords)
    location = random.choice(locations) + keyword
    duration = random.choice(durations)
    price = random.choice(prices)
    description = random.choice(descriptions)

    # 📅 랜덤 날짜 (2024~2025년, 1~12월, 1~28일)
    year = random.choice(["2024", "2025"])
    month = f"{random.randint(1, 12):02d}"
    day = f"{random.randint(1, 28):02d}"

    # 파일명에서 사용할 수 없는 문자 제거
    safe_description = re.sub(r'[<>:"/\\|?*]', '', description)

    # 📄 문서 작성
    doc.add_heading("입찰공고", level=1)
    doc.add_paragraph(f"사업명 : {keyword} 기반 {description}")
    doc.add_paragraph(f"회사명 : {company}")
    doc.add_paragraph(f"사업기간 : {year}년 {int(month)}월 ~ {year}년 {int(month)+1 if int(month) < 12 else 12}월")
    doc.add_paragraph(f"설치위치 : {location}")
    doc.add_paragraph(f"계약금액 : {price:,}원 (VAT 포함)")
    doc.add_paragraph(f"사업내용 : 본 사업은 {keyword} 기술을 활용하여 {description}을 수행하고자 함.")

    # 📁 파일명에 날짜 포함
    #filename = f"./docs/{year}{month}{day}_{company}_{safe_description}{i:02d}.docx"
    filename = f"./docs/{year}{month}{day}_{company}_{safe_description}.docx"
    doc.save(filename)

print("연/월/일이 포함된 mock 문서 100개 './docs' 폴더에 생성")