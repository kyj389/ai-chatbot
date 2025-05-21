import os
import time
import streamlit as st
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from docx import Document as DocxLoader
import fitz  # pymupdf

def read_docx(path: str):
    doc = DocxLoader(path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text, metadata={"source": path.split('/')[-1]})]

def read_pdf(path: str):
    # loader = PyPDFLoader(file_path)
    # documents = loader.load()
    # return documents

    doc = fitz.open(path)
    all_text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)  # 페이지 로드
        text = page.get_text()          # 텍스트 추출
        all_text += f"\n{text}"

    doc.close()
    return [Document(page_content=all_text, metadata={"source": path.split('/')[-1]})]

def get_document(file_path: str):
    if file_path.endswith(".docx"):
        return read_docx(file_path)
    elif file_path.endswith(".pdf"):
        return read_pdf(file_path)
    else:
        print(f"지원하지 않는 파일 형식입니다.{file_path.split('/')[-1]}")

def create_chroma_db(dir_path):
    
    documents = []
    # collection = chromadb.Client().create_collection(name="docs")

    for file_name in os.listdir(dir_path):
        print(f"file_name : {file_name}")
        file_path = os.path.join(dir_path, file_name)
        documents += get_document(file_path)
        

    # 3. 문서 분할 (chunking)
    # text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    text_splitter = CharacterTextSplitter(chunk_size=1500, chunk_overlap=150)
    split_docs = text_splitter.split_documents(documents)


    # 4. 임베딩 모델 준비 (OpenAI)
    embedding = OpenAIEmbeddings()

    # 5. Chroma 벡터 DB에 저장
    vectordb = Chroma.from_documents(split_docs, embedding, persist_directory="./chroma_db")
    vectordb.persist()

    return vectordb

def load_chroma_db():
    embeddings = OpenAIEmbeddings()
    vectordb = Chroma(
        persist_directory="./chroma_db",
        embedding_function=embeddings
    )
    return vectordb









class Chroma2:

    def create_langchain():
        
        start_time = time.time()

        # print(f"OPENAI_API_KEY={os.environ["OPENAI_API_KEY"]}")
        dir_path = "./Documents"

        # 최초 1회만 임베딩 및 벡터스토어 초기화
        if "vectordb" not in st.session_state:
            vectordb = load_chroma_db()
            print(len(vectordb))

            if len(vectordb) == 0:
                # 벡터DB 초기화
                print("✅ 벡터DB 파일 생성")
                vectordb = create_chroma_db(dir_path)
           
            st.session_state.vectordb = vectordb
            print("✅ 벡터DB 파일 불러오기")
        else:
            print("📦 벡터DB 세션 불러오기")
            

        vectordb = st.session_state.vectordb
        print(len(vectordb))
        elapsed = time.time() - start_time
        print(f"⏱️ 벡터DB 수행 : {elapsed:.4f}초")


        # 6. Retriever 생성
        retriever = vectordb.as_retriever(search_kwargs={"k": 3})

        custom_prompt = PromptTemplate(
            input_variables=["context", "question"],
            template=
            "당신은 신한금융그룹의 ICT 사업정보를 분석해주는 전문 챗봇입니다."
            "다음 문서를 기반으로 사실에 근거한 정보를 명확하고 친절하게 제공합니다."
            "문서:"
            "{context}"
            "질문:"
            "{question}"
        )
        # print(custom_prompt)        

        # 7. QA 체인 생성 (Retriever + ChatGPT)
        qa_chain = RetrievalQA.from_chain_type(
            llm=ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.2),
            # llm=ChatOpenAI(model_name="gpt-4o-mini", temperature=0.2),
            # llm=ChatOpenAI(model_name="gpt-4", temperature=0.2),
            chain_type="stuff", 
            retriever=retriever,
            return_source_documents=True,
            chain_type_kwargs={"prompt": custom_prompt}
            # chain_type="map_reduce"
            # chain_type_kwargs={"question_prompt": chat_prompt}
        )

        return qa_chain


# 8. 질의
# query = "신한은행의 최근 신사업은?"
# query = "베트남은행 사업 입찰 자격에 대해서 알려줘"
# langchain = Chroma2.create_langchain()
# result = langchain(query)

# print("\n📌 답변:")
# print(result["result"])

# print("\n📚 관련 문서:")
# for doc in result["source_documents"]:
#     print(f"- {doc.metadata['source']}")